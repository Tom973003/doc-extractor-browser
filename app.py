import streamlit as st
import fitz  # PyMuPDF
import docx
import zipfile
import re
import io
from PIL import Image

# -------------------------------------------------
# Page setup
# -------------------------------------------------
st.set_page_config(page_title="RFQ → Important Bits")
st.title("📄 RFQ → Important Bits")

uploaded_file = st.file_uploader(
    "Upload RFQ document (PDF or DOCX)",
    type=["pdf", "docx"]
)

# -------------------------------------------------
# Helpers
# -------------------------------------------------
def clean(text):
    return re.sub(r"\s+", " ", text.lower()).strip()

def safe_image(img):
    try:
        img = img.convert("RGB")
        if img.width > 2000:
            ratio = 2000 / img.width
            img = img.resize(
                (2000, int(img.height * ratio)),
                Image.LANCZOS
            )
        return img
    except Exception:
        return None

# -------------------------------------------------
# PDF extraction
# -------------------------------------------------
def extract_pdf_text_and_images(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    images = []

    for page in doc:
        text += page.get_text()
        pix = page.get_pixmap(dpi=120)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        img = safe_image(img)
        if img:
            images.append(img)

    return text, images

# -------------------------------------------------
# DOCX extraction
# -------------------------------------------------
def extract_docx_text_tables_images(file):
    doc = docx.Document(file)
    text = ""
    table_rows = []
    images = []

    for p in doc.paragraphs:
        text += p.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                table_rows.append((clean(cells[0]), cells[1]))

    with zipfile.ZipFile(file) as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                img = Image.open(io.BytesIO(z.read(name)))
                img = safe_image(img)
                if img:
                    images.append(img)

    return text, table_rows, images

# -------------------------------------------------
# Field extraction
# -------------------------------------------------
FIELDS = {
    "Project Title": ["project title"],
    "Description of Work": ["description of work"],
    "RFQ Close": ["rfq close"],
    "Proposed Start Date": ["proposed start date", "start date"],
    "Proposed Completion Date": ["proposed completion date", "completion date"],
    "Site Location": ["site location"],
    "LRD": ["lrd"],
    "Inc": ["inc"],
    "Tas": ["tas"],
    "Practical Work": ["practical work", "scope of works"],
}

def extract_fields(text, table_rows):
    results = {}
    text_clean = clean(text)

    for field, keys in FIELDS.items():
        value = "Not found"

        for k, v in table_rows:
            for key in keys:
                if key in k:
                    value = v
                    break
            if value != "Not found":
                break

        if value == "Not found":
            for key in keys:
                if key in text_clean:
                    idx = text_clean.find(key)
                    value = text[idx:idx + 300].replace("\n", " ").strip()
                    break

        results[field] = value

    return results

# -------------------------------------------------
# Main
# -------------------------------------------------
if uploaded_file:
    text = ""
    table_rows = []
    images = []

    if uploaded_file.type == "application/pdf":
        text, images = extract_pdf_text_and_images(uploaded_file)
    else:
        text, table_rows, images = extract_docx_text_tables_images(uploaded_file)

    fields = extract_fields(text, table_rows)

    st.subheader("📌 Extracted Fields")
    for k, v in fields.items():
        st.markdown(f"**{k}:** {v}")
        st.divider()

    st.subheader("🖼️ Extracted Visuals")
    if images:
        for i, img in enumerate(images):
            st.image(img, caption=f"Image / Page {i + 1}", use_container_width=True)
    else:
        st.info("No images detected.")
